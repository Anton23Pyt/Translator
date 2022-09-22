# -*- coding: utf-8 -*-
import re
import docx
from selenium import webdriver
from webdriver_manager.firefox import GeckoDriverManager
import time as ti
from random import randint
from selenium.webdriver.firefox.service import Service

def get(driver,c1):
    ti.sleep(randint(30, 60))
    d = driver.find_element_by_css_selector("textarea.lmt__source_textarea")
    d.send_keys(c1)
    ti.sleep(3)
    x = driver.find_element_by_css_selector("textarea.lmt__target_textarea")
    c1 = x.get_attribute("value")
    while c1 == '':
        x = driver.find_element_by_css_selector("textarea.lmt__target_textarea")
        c1 = x.get_attribute("value")
        ti.sleep(randint(1, 3))
    ti.sleep(2)
    c1 = x.get_attribute("value")
    return c1



path=input("Введите распложение файла:")
m=input('Введите количество символов:')
m=int(m)
doc = docx.Document(path)
text = []
for paragraph in doc.paragraphs:
    text.append(paragraph.text)
r='\n'.join(text)
driver = webdriver.Firefox(service=Service(GeckoDriverManager().install()))
driver.get("https://www.deepl.com/translator")
input("введи ну хоть шото")
t3=[]
for i in r:
    t3.append(i)
o3=False
e3=0
w3=[]
y3=False
b='@'
result1=[]
result2=[]
just=True
just1=0
fu=0
while e3<len(t3):
    if e3==len(t3)-1:
        y3=True
    try:
        if t3[e3] == b and t3[e3 + 1] == b and t3[e3 + 2] == b and o3 == False:
            o3 = True
            e3 = e3 + 3
    except:
        otyi=1
    try:
        if t3[e3] == b and t3[e3 + 1] == b and t3[e3 + 2] == b and o3 == True:
            o3 = False
            e3 = e3 - 2
            y3 = True
    except:
        orfks=1
    if o3==True:
        w3.append(t3[e3])
    e3=e3+1
    if y3==True:
        f=w3[0]
        for i in range(1,len(w3)):
            f=f+w3[i]
        w3=[]
        y3=False
        result1.append(f)
        g = []
        g = f.split("\n")
        t = 0
        c = 0
        o = False
        a = []
        while t < len(g):
            w = 0
            for h in g[t]:
                w = w + 1
            if o == False:
                c1 = g[t]
                o = True
            c = c + w
            if c < m:
                if c1 != g[t]:
                    c1 = c1 + '\n' + '\n' + g[t]
            if c >= m:
                if c > w:
                    c = 0
                    c1=get(driver, c1)
                    a.append(c1)
                    ti.sleep(5)
                    try:
                        d = driver.find_element_by_css_selector("button.button--lPYGv").click()
                    except:
                        e3=len(t3)+10
                        t=len(g)+10
                    o = False
                    t = t - 1
                if c == w:
                    p = []
                    for s in re.split(r'(?<=[.!?…]) ', g[t]):
                        p.append(s)
                    t1 = 0
                    c2 = 0
                    a2 = []
                    o2 = False
                    while t1 < len(p):
                        if o2 == False:
                            c3 = p[t1]
                            o2 = True
                        w1 = 0
                        for h in p[t]:
                            w1 = w1 + 1
                        c2 = c2 + w1
                        if c2 < m:
                            if c3 != p[t1]:
                                c3 = c3 + p[t1]
                        if c2 >= m:
                            c3=get(driver, c3)
                            a2.append(c3)
                            ti.sleep(5)
                            try:
                                d = driver.find_element_by_css_selector("button.button--lPYGv").click()
                            except:
                                fu=t1
                                e3 = len(t3) + 10
                                t = len(g) + 10
                                t1=len(p)+10
                                just=False
                            if just==False and  fu<len(p)-1:
                                just1=1
                            t1 = t1 - 1
                            c2 = 0
                            o2 = False
                        if t1 == len(p) - 1 and c2 != 0:
                            c3=get(driver, c3)
                            a2.append(c3)
                            ti.sleep(5)
                            try:
                                d = driver.find_element_by_css_selector("button.button--lPYGv").click()
                            except:
                                e3 = len(t3) + 10
                                t = len(g) + 10
                                t1=len(p)+10
                        t1 = t1 + 1
                    f1 = a2[0]
                    for i in range(1, len(a2)):
                        f1 = f1 + a2[i]
                    a.append(f1)
                    o = False
                    c = 0
            if t == len(g) - 1 and c != 0:
                ti.sleep(randint(30, 60))
                c1=get(driver, c1)
                a.append(c1)
                ti.sleep(5)
                try:
                    d = driver.find_element_by_css_selector("button.button--lPYGv").click()
                except:
                    e3 = len(t3) + 10
                    t = len(g) + 10
            t = t + 1
        c1 = a[0]
        for i in range(1, len(a)):
            c1 = c1 + '\n' + a[i]
        result2.append(c1)
driver.close()
assert len(result1)==len(result2)
if just1==1:
    result1.pop(len(result1)-1)
    result2.pop(len(result2) - 1)
rok=result1[0]+'\n'+result2[0]+'\n'+'=================================================='
for i in range(1,len(result1)):
    rok=rok+'\n'+result1[i]+'\n'+result2[i]+'\n'+'=================================================='
doc2 = docx.Document()
par=doc2.add_paragraph(rok)
doc2.save(path)


